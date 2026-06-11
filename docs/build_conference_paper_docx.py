#!/usr/bin/env python3
"""Build docs/CONFERENCE_PAPER.docx from docs/CONFERENCE_PAPER.md in MLA dress.

Run from the repository root:

    python docs/build_conference_paper_docx.py

This is the conference edition of the project paper ("Improving the Experience
of Observing Archived Data"). The pipeline is shared with the technical
reference edition by importing build_technical_paper_docx and reusing its
python-docx MLA post-processing, so the two editions cannot drift in dress.
The one reader difference is +implicit_figures, which turns each image's alt
text into a real captioned figure (Figure 1 through Figure 10) instead of a
bare image.
"""

import subprocess
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_technical_paper_docx as base

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "CONFERENCE_PAPER.md"
DOCX = ROOT / "docs" / "CONFERENCE_PAPER.docx"
EXPECTED_FIGURES = 10


def run_pandoc():
    cmd = [
        base.find_pandoc(), str(MD), "-o", str(DOCX),
        "--from", "gfm+yaml_metadata_block+implicit_figures",
        "--resource-path", str(ROOT / "docs"),
    ]
    subprocess.run(cmd, check=True)
    print("pandoc:", " ".join(cmd[1:]))


def verify():
    base.verify()  # Calibri, running head with PAGE field, margins, Works Cited
    doc = Document(str(DOCX))

    figures = len(doc.inline_shapes)
    assert figures == EXPECTED_FIGURES, f"expected {EXPECTED_FIGURES} figures, found {figures}"

    captions = [p.text for p in doc.paragraphs
                if "Caption" in p.style.name and p.text.strip()]
    assert len(captions) >= EXPECTED_FIGURES, (
        f"expected {EXPECTED_FIGURES} figure captions, found {len(captions)}")

    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    joined = "\n".join(texts)
    assert "—" not in joined, "em dash (U+2014) found in document text"

    words = sum(len(t.split()) for t in texts)
    print(f"verified: {figures} figures, {len(captions)} captions, "
          f"no em dashes, ~{words} words including tables")


if __name__ == "__main__":
    run_pandoc()
    base.MD = MD
    base.DOCX = DOCX
    base.post_process()
    verify()
