#!/usr/bin/env python3
"""Build docs/TECHNICAL_PAPER.docx from docs/TECHNICAL_PAPER.md in MLA dress.

Run from the repository root:

    python docs/build_technical_paper_docx.py

Pipeline: pandoc converts the Markdown (the exact invocation is below), then
python-docx applies the formatting the paper's audience requires: Calibri
throughout, 12 pt double-spaced body prose with half-inch first-line indents,
one-inch margins, an MLA running head (first author's surname + page number,
top right, every page), black headings, single-spaced tables/code/captions,
and a hanging-indent Works Cited. The script is deterministic; re-running it
reproduces the same document from the same source.
"""

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "TECHNICAL_PAPER.md"
DOCX = ROOT / "docs" / "TECHNICAL_PAPER.docx"
RUNNING_HEAD_NAME = "O'Hara"  # MLA: first listed author's surname

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"


def find_pandoc():
    p = shutil.which("pandoc")
    if p:
        return p
    candidates = [
        Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe",
        Path("C:/Program Files/Pandoc/pandoc.exe"),
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    sys.exit("pandoc not found; install pandoc 3.x first")


def run_pandoc():
    cmd = [
        find_pandoc(), str(MD), "-o", str(DOCX),
        "--from", "gfm+yaml_metadata_block",
        "--resource-path", str(ROOT / "docs"),
    ]
    subprocess.run(cmd, check=True)
    print("pandoc:", " ".join(cmd[1:]))


def set_style_font(style, name, size=None):
    try:
        style.font.name = name
        if size is not None:
            style.font.size = size
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            fonts.set(qn(attr), name)
    except (AttributeError, ValueError):
        pass


def style_or_none(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def add_running_head(doc):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(RUNNING_HEAD_NAME + " ")
    run.font.name = BODY_FONT
    run.font.size = Pt(12)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld_r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(attr), BODY_FONT)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # half-points
    rpr.append(fonts)
    rpr.append(sz)
    fld_r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    fld_r.append(t)
    fld.append(fld_r)
    p._p.append(fld)


def post_process():
    doc = Document(str(DOCX))

    # Every named style gets Calibri; code styles get Consolas.
    for style in doc.styles:
        if style.name in ("Source Code", "Verbatim Char", "HTML Code"):
            set_style_font(style, CODE_FONT, Pt(9.5))
        else:
            set_style_font(style, BODY_FONT)

    normal = style_or_none(doc, "Normal")
    if normal:
        normal.font.size = Pt(12)

    # MLA body prose: double-spaced, half-inch first-line indent.
    for name in ("Body Text", "First Paragraph"):
        s = style_or_none(doc, name)
        if s:
            pf = s.paragraph_format
            pf.line_spacing = 2.0
            pf.first_line_indent = Inches(0.5)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    # Title block: centered plain title, left-aligned author/date lines.
    for name, align, size, italic in (
        ("Title", WD_ALIGN_PARAGRAPH.CENTER, Pt(16), False),
        ("Subtitle", WD_ALIGN_PARAGRAPH.CENTER, Pt(13), True),
        ("Author", WD_ALIGN_PARAGRAPH.LEFT, Pt(12), False),
        ("Date", WD_ALIGN_PARAGRAPH.LEFT, Pt(12), False),
    ):
        s = style_or_none(doc, name)
        if s:
            s.paragraph_format.alignment = align
            s.paragraph_format.line_spacing = 2.0
            s.font.size = size
            s.font.italic = italic
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.font.bold = name == "Title"

    # Headings: black, Calibri (size left at pandoc defaults for scannability).
    for i in range(1, 7):
        s = style_or_none(doc, f"Heading {i}")
        if s:
            s.font.color.rgb = RGBColor(0, 0, 0)

    # Captions and figures: single-spaced, centered, no indent.
    for name in ("Image Caption", "Table Caption", "Captioned Figure", "Figure"):
        s = style_or_none(doc, name)
        if s:
            pf = s.paragraph_format
            pf.line_spacing = 1.0
            pf.first_line_indent = Inches(0)
            if name in ("Captioned Figure", "Figure"):
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sc = style_or_none(doc, "Source Code")
    if sc:
        sc.paragraph_format.line_spacing = 1.0
        sc.paragraph_format.first_line_indent = Inches(0)

    compact = style_or_none(doc, "Compact")
    if compact:
        compact.paragraph_format.line_spacing = 1.5
        compact.paragraph_format.first_line_indent = Inches(0)

    # One-inch margins, MLA.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_running_head(doc)

    # Tables stay single-spaced regardless of body-style inheritance.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.first_line_indent = Inches(0)

    # Works Cited: hanging indent, double-spaced, left-aligned.
    in_wc = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            in_wc = p.text.strip() == "Works Cited"
            continue
        if in_wc and p.text.strip():
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            pf.line_spacing = 2.0
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(str(DOCX))


def verify():
    doc = Document(str(DOCX))
    assert doc.styles["Normal"].font.name == BODY_FONT
    header_xml = doc.sections[0].header._element.xml
    assert 'w:instr="PAGE"' in header_xml and RUNNING_HEAD_NAME in header_xml
    assert doc.sections[0].left_margin == Inches(1)
    wc = [p for p in doc.paragraphs if p.text.strip() == "Works Cited"]
    assert wc, "Works Cited heading missing"
    print("verified: Calibri Normal, MLA running head with PAGE field, "
          "1-inch margins, Works Cited present")
    print("size:", DOCX.stat().st_size, "bytes")


if __name__ == "__main__":
    run_pandoc()
    post_process()
    verify()
